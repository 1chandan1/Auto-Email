import base64
import json
import locale
import msvcrt
import os
import pickle
import random
import re
import shutil
import sys
from datetime import datetime
from dateutil.relativedelta import relativedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from time import sleep
from updater import check_for_updates

import gspread
from docx import Document
from docx2pdf import convert
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from gspread.exceptions import SpreadsheetNotFound
from unidecode import unidecode
from dotenv import load_dotenv


class GoogleServices:

    SCOPES = [
        "https://www.googleapis.com/auth/gmail.compose",
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive"
    ]
    INTERN_SHEET_KEY = "1cveFT3BvSJ9d-PvBdyGvZ7Fd_0XYDlNafOpTXUiz_eI"

    def __init__(self):
        self.creds = None
        self.gmail_service = None
        self.drive_service = None
        self.sheet_service = None
        self.gmail = None
        self.sender_name = None
        self.phone = None
        self.gender = None
        self.login()
        self.signature = self.get_signature()
        
    def login(self):
        self.google_auth()

        clear_display()
        while True:
            print("\n")
            print_center(f"Last Logged-in Account : {self.email}")
            print("\n")
            print_center("Do you want to use it (y/n) : ")
            while True:
                if msvcrt.kbhit():
                    choice = msvcrt.getch().decode('utf-8').lower()
                    if choice == "y" or choice == "n":
                        print("\nLoading...")
                        break
                sleep(0.1)
            while msvcrt.kbhit():
                msvcrt.getch()
            if choice == "n":
                self.google_auth(True)
            if self.gmail in self.interns.keys():
                self.sender_name = self.interns[self.gmail]["Name"]
                self.phone = self.interns[self.gmail]["Phone"]
                self.gender = self.interns[self.gmail]["Gender"]
                break
            else:
                clear_display()
                print("\n\n\nInvalid Sender-Email")
                print("Use any of these Account :")
                for i in self.interns:
                    print(f"  {i}   ({self.interns[i]['Email']})")
                print()

    def google_auth(self, new=False):
        """This will help to create service for the object"""
        if new:
            self.creds = None
        else:
            # The file secret_token.pickle stores the user's access and refresh tokens
            if os.path.exists('secret_token.pickle'):
                with open('secret_token.pickle', 'rb') as token:
                    self.creds = pickle.load(token)

        # If there are no (valid) credentials available, let the user log in
        if not self.creds or not self.creds.valid:
            try:
                self.creds.refresh(Request())
            except:
                # Use the JSON file containing your OAuth2 credentials
                flow = InstalledAppFlow.from_client_config(client_secret_info, self.SCOPES)
                self.creds = flow.run_local_server(port=0)

            # Save the credentials for the next run
            with open('secret_token.pickle', 'wb') as token:
                pickle.dump(self.creds, token)

        self.gmail_service = build('gmail', 'v1', credentials=self.creds)
        self.drive_service = build('drive', 'v3', credentials=self.creds)
        self.sheet_service = build('sheets', 'v4', credentials=self.creds)
        profile = self.gmail_service.users().getProfile(userId='me').execute()
        self.gmail = profile['emailAddress']
        self.interns = self.get_interns_info()
        self.email = self.get_user_email()

    def get_interns_info(self):
        gc = gspread.authorize(self.creds)
        spreadsheet = gc.open_by_key(self.INTERN_SHEET_KEY) 
        worksheet = spreadsheet.get_worksheet(0)

        interns_data = worksheet.get_all_records()

        interns_dict = {}
        for intern in interns_data:
            interns_dict[intern['Gmail']] = {
                'Name': intern['Name'],
                'Phone': intern['Phone'],
                'Email': intern['Email'],
                'Gender': intern['Gender']
            }
        return interns_dict
    
    def get_user_email(self):
        try:
            email =  self.interns[self.gmail]["Email"]
            if email == "":
                return self.gmail
            return email
        except:
            return self.gmail
        
    def get_signature(self):
        return f'''
            <div>
            <b><p>{self.sender_name}</p>
            <p>{"Chargé d'affaires" if self.gender == "M" else "Chargée d'affaires" }</p>
            <p>{self.phone}</p>
            <p>Prendre rendez-vous directement ici: <a href="https://calendly.com/klero-genealogie1/30min">Calendly</a></p>
            <p>p/o Jean-Michel Branche Gérant Klero Généalogie, <a href="https://www.avocatparis.org/annuaire/avocat?cnbf=2nWRmtPEdadV+CSnc4jCEA==">Avocat honoraire</a></p>
            <p>295 Rue Saint-Jacques, 75005 Paris (Siège social)</p>
            <p>TVA intracommunautaire: FR49947791927</p>
            <p>01 73 71 13 63</p></b>
            <img src="https://drive.google.com/uc?export=view&id=1Z4E-R3MxulRpkMQa4qNAV2LoFrKUxTgo" width="125" height="75">
            </div>'''


def get_filled_row_values(worksheet , row_number):
    merged_ranges = worksheet.spreadsheet.fetch_sheet_metadata()['sheets'][0]['merges']
    row_values = worksheet.row_values(row_number)
    
    # Determine the maximum column index in merged_ranges
    max_col_index = max(merge['endColumnIndex'] for merge in merged_ranges)
    # Extend row_values to cover the entire range of columns
    row_values.extend([''] * (max_col_index - len(row_values)))
    
    # Process each merged range
    for merge in merged_ranges:
        # If the merge affects the row in question
        if merge['startRowIndex'] < row_number <= merge['endRowIndex']:
            # Get the value from the first cell of the merged range
            start_col_index = merge['startColumnIndex']
            end_col_index = merge['endColumnIndex']
            merged_value = worksheet.cell(row_number, start_col_index + 1).value

            # Apply this value to all cells in the merged range within the row
            for col_index in range(start_col_index, end_col_index):
                row_values[col_index] = merged_value

    return row_values

def find_column_index(primary_header_row, secondary_header_row, target_secondary_heading, target_primary_heading=None):
    # Iterate over both header rows simultaneously
    for index, (primary, secondary) in enumerate(zip(primary_header_row, secondary_header_row)):
        # Check if the current column's secondary header matches the target
        if secondary == target_secondary_heading:
            # If a primary header is specified, check for a match; otherwise, return the current index
            if target_primary_heading is None or primary == target_primary_heading:
                return index

    # Return None if no matching column is found
    return None

    
def get_row_by_name(first_name : str ,last_name : str):
    pattern = r'[ ,\-\n]'
    first_name = unidecode(re.sub(pattern, '', first_name)).lower()
    last_name = unidecode(re.sub(pattern, '', last_name)).lower()
    all_rows = notary_worksheet.get_all_values()
    for index, row in enumerate(all_rows, start=1):
        if first_name == unidecode(re.sub(pattern, '', row[1])).lower() and last_name == unidecode(re.sub(pattern, '', row[2])).lower():
            return index, row
    return None, None


def print_center(text):
    terminal_width = shutil.get_terminal_size().columns
    padding = (terminal_width - len(text)) // 2
    print(" " * padding + text)


def send_email(message: MIMEMultipart):
    try:
        status = user.gmail_service.users().messages().send(
            userId=user.gmail, body=message).execute()
        if status:
            print("\nEmail sent successfully.")
            sleep(2)
            return status
    except Exception as e:
        print(f"Error sending email: {e}")
        return None


def create_draft(message: MIMEMultipart):
    try:
        status = user.gmail_service.users().drafts().create(
            userId=user.gmail, body={'message': message}).execute()
        if status:
            return status
    except Exception as e:
        print(f"Error creating draft: {e}")


def create_notary_message(sender: str, to: str, person_full_name: str, person_last_name: str, notary_last_name: str, person_don: str):
    message = MIMEMultipart()
    message['From'] = f"Klero Genealogy <{sender}>"
    message['To'] = to
    message['Subject'] = f'Succession {person_last_name} - Demande de mise en relation'
    message_html = f'''
            <p>À l'attention de Maître {notary_last_name}</p>
            <p>Maître,</p>
            <p>Dans le cadre de notre activité, nous avons développé une nouvelle prestation dédiée à la recherche de bénéficiaires d'actifs non réclamés.</p>
            <p>Votre étude s'est chargée de régler la succession de {person_full_name} dont l'acte de notoriété a été établi le {person_don}. Toutefois, d'après nos recherches, il est probable que des fonds non connus par la famille n'aient pas été transmis aux héritiers. Nous n'en connaissons, pour le moment, ni le support (par ex. compte bancaire, assurance vie, plan épargne retraite, épargne salariale) ni le montant. Ces montants sont généralement limités (en moyenne 500€ sur les 500 dossiers traités à l'étude l'an dernier), mais il nous arrive de retrouver des montants plus conséquents.</p>
            <p>Ne pouvant nous mandater nous-mêmes, nous avons besoin de rentrer en contact avec les héritiers afin de proposer notre prestation pour obtenir les informations précitées et débloquer lesdits fonds. Ainsi, <b>pouvez-vous transmettre mes coordonnées à l'un des héritiers afin que ce dernier puisse revenir vers moi pour de plus amples renseignements ?</b></p>
            <p>À titre informatif, sachez que :</p>
            <ul>
            <li>Si la succession est toujours ouverte, les fonds débloqués seront réintégrés déduits de nos honoraires</li>
            <li>Si la succession est clôturée, les fonds seront directement reversés aux héritiers, et nous vous aviserons si le montant de ces derniers pourrait avoir un impact sur les droits.</li>
            </ul>
            <p>Nous tenons à souligner, conformément à la loi Eckert du 13 juin 2014, l'importance cruciale de l'information complète aux héritiers pour le règlement équitable des successions. Vous remerciant par avance de votre concours.</p>
            <p>Bien cordialement,</p>
        '''
    message.attach(MIMEText(message_html + user.signature, 'html'))

    return {'raw': base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')}


def create_client_message(sender: str, to: str, person_full_name: str, amount_found_by_us: str, amount_with_tex: str, amount_after_fee: str):
    message = MIMEMultipart()
    message['From'] = f"Klero Genealogy <{sender}>"
    message['To'] = to
    message['Subject'] = f'Retour sur actifs débloqués - {person_full_name}'
    message_html = f'''
            <p>Bonjour,</p>
            <p>Je reviens vers vous concernant les actifs au nom de {person_full_name}.</p>
            <p>Suite au retour du notaire, je vous informe que les fonds débloqués s'élèvent à {amount_found_by_us} (voir pièce jointe).</p>
            <p>Conformément au contrat précédemment signé, nos honoraires sont de {amount_with_tex} TTC de sorte que la somme vous revenant est de {amount_after_fee}.</p>
            <p>Au regard des articles 11 et 12 du contrat, deux options s'offrent à vous :</p>
            <ul>
                <li>Récupérer ces fonds auquel cas il convient de nous faire parvenir votre RIB par mail ou par courrier. Nous procéderons à une vérification avant tout envoi des fonds</li>
                <li>Faire don de la somme à une association de notre choix. Dans cette hypothèse, il convient impérativement  de nous donner votre accord par écrit en réponse à ce mail</li>
            </ul>
            <p>Nous vous informons qu'en cas de non retour de votre part sur votre choix dans un délai de deux mois à compter de la réception de ce mail, nous verserons automatiquement les fonds à une association.</p>
            <p>Conformément au RGPD, nous vous informons que nous supprimerons à la clôture du dossier de manière sécurisée la copie de votre RIB.</p>
            <p>Je reste à votre disposition pour répondre à d'éventuelles questions par téléphone au 07.45.25.93.99.</p>
            <p>Vous remerciant par avance pour votre retour,</p>
            <p>Bien cordialement,</p>
        '''
    message.attach(MIMEText(message_html + user.signature, 'html'))

    return {'raw': base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')}


def create_facture_message(sender: str, to: str, person_full_name: str):
    message = MIMEMultipart()
    if sender:
        message['From'] = f"Klero Genealogy <{sender}>"
    message['To'] = to
    message['Subject'] = f'Clôture dossier {person_full_name}'
    message_html = f'''
            <p>Bonjour,</p>
            <p>Je reviens vers vous concernant les actifs au nom de {person_full_name}</p>
            <p>Je vous informe que l'étude vous a transmis les fonds vous revenant déduit de nos honoraires. Pour rappel, en cas de pluralité d'héritiers, vous vous êtes engagés, en signant le contrat, à faire le partage desdits fonds entre les différents héritiers.</p>
            <p>Notre mission étant désormais terminée, je vous remercie de votre confiance et vous invite à laisser un avis sur la page Google de LD Généalogie. En effet, comme vous l'étiez probablement lors de notre premier échange, les bénéficiaires sont souvent méfiants vis-à-vis de notre démarche. Ainsi, votre témoignage pourra les rassurer et nous permettre de débloquer et restituer d'avantage de fonds.</p>
            <p>Vous souhaitant une bonne continuation,</p>
            <p>Bien cordialement,</p>
        '''
    message.attach(MIMEText(message_html + user.signature, 'html'))

    return {'raw': base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')}


def create_facture_files(name, facture_number, ht, tva, tcc, paid_date):
    if not os.path.exists("Invoice"):
        os.makedirs("Invoice")
    doc = Document(resource_path("template.docx"))
    date = datetime.now().date().strftime('%d %B %Y')
    text_replacements = {
        "(DATE)": date,
        "(B)": name,
        "(Q)": facture_number,
        "(R)": ht,
        "(S)": tva,
        "(T)": tcc,
        "(W)": paid_date,
    }

    def modify_run(run):
        for find_text, replace_text in text_replacements.items():
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)

    # Replace text in Header
    header = doc.sections[0].header
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            modify_run(run)

    # Replace text in Body
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            modify_run(run)

    # Replace text in Table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        modify_run(run)

    # Save the modified document
    doc.save(f"Invoice/{facture_number} {name}.docx")
    convert(f"Invoice/{facture_number} {name}.docx")



def send_notary_emails(spreadsheet: gspread.Spreadsheet):
    worksheet = spreadsheet.get_worksheet(0)
    all_values = worksheet.get_all_values()
    for index, row in enumerate(all_values, start=1):
        if row[10] == "à envoyer":
            notary_email = str(row[8]).split("\n")[0]
            person_full_name = str(row[0]).strip()
            words = person_full_name.split()
            person_last_name = " ".join(
                [word for word in words if word.isupper()])
            if not person_last_name.strip():    
                continue
            notary_full_name = str(row[5]).strip()
            words = notary_full_name.split()
            notary_last_name = " ".join(
                [word for word in words if word.isupper()])
            notary_first_name = notary_full_name.replace(
                notary_last_name, "").strip()
            if not notary_last_name.strip():
                continue
            person_don = row[4]
            notary_sheet_index, notary_sheet_row = get_row_by_name(
                notary_first_name, notary_last_name)
            if not notary_sheet_index:
                worksheet.update_acell(f"L{index}", "New Notary added")
                first_col = notary_worksheet.col_values(2)  # Get all values in the first column

                notary_sheet_row = ["", notary_first_name, notary_last_name, "", "","", row[6], row[7], row[9], row[8], "Not contacted", ""]
                notary_sheet_index = len(first_col) + 1

                notary_worksheet.insert_row(
                    notary_sheet_row, index=notary_sheet_index, inherit_from_before=True)

            if notary_sheet_row[10] == "Not cooperating":
                worksheet.update_acell(f"L{index}", "Not cooperating")
                continue
            contact_date = notary_sheet_row[11]
            clear_display()
            print("\n")
            print_center(
                f"-------------------  Account : {user.email}  -------------------")
            print()
            print_center("Notary Email")
            print()
            print_center(f"Google Sheet : {spreadsheet.title}")
            print()
            print_center("Sending All Emails\n\n")
            print(f"Index-File Row    :    {notary_sheet_index}")
            print(f"Contact Date      :    {contact_date}\n")
            print(f"Target Sheet Row  :    {index + 1}\n")
            print(f"Person Name       :    {person_full_name}")
            print(f"Person Last Name  :    {person_last_name}\n")
            print(f"Notary Name       :    {notary_full_name}")
            print(f"Notary Last Name  :    {notary_last_name}\n")
            print(f"DON               :    {person_don}")
            print(f"To                :    {notary_email}\n")
            notary_worksheet.update_acell(f"J{notary_sheet_index}", notary_email)
            
            if notary_sheet_row[10] == "Not contacted":
                countdown("Sending Email in", random.randint(120, 180))
                print("\nSending Email...")
                for _ in range(3):
                    message = create_notary_message(user.email, notary_email, person_full_name, person_last_name, notary_last_name, person_don)
                    status = send_email(message)
                    if status:
                        worksheet.update_acell(f"K{index}", "envoyé")
                        today_date = datetime.now().date().strftime("%d/%m/%Y")
                        all_row_with_same_email = [index for index, sublist in enumerate(notary_worksheet.get_all_values(),start=1) if notary_email in sublist]
                        for row_index in all_row_with_same_email:
                            notary_worksheet.update_acell(f"L{row_index}", today_date)
                            notary_worksheet.update_acell(f"K{row_index}", "Contacted / pending answer")
                            notary_worksheet.update_acell(f"O{row_index}", person_full_name)
                        print("\nSuccess")
                        break
                    else:
                        countdown("Trying to send email again", 5)
                        pass
            elif notary_sheet_row[10] in ("Contacted / pending answer","Cooperating"):
                print("Scheduling Email")
                sleep(2)
                previous_scheduled_date = notary_sheet_row[11]
                all_scheduled_data = scheduling_worksheet.get_all_values()
                for scheduled_data in all_scheduled_data[::-1]:
                    if notary_email in scheduled_data and scheduled_data[3]=="Scheduled":
                        previous_scheduled_date = scheduled_data[7]
                        break
                new_date = (datetime.strptime(previous_scheduled_date, "%d/%m/%Y") + relativedelta(months=+2)).strftime("%d/%m/%Y")
                scheduling_worksheet.append_row([notary_first_name,notary_last_name,None,"Scheduled",person_full_name,user.email,notary_email,new_date])
                worksheet.update_acell(f"L{index}", f"Scheduled on {new_date}")
                worksheet.update_acell(f"K{index}", "draft")
                print(f"Scheduled on {new_date}")
                    


def clear_display():
    os.system('cls' if os.name == 'nt' else 'clear')


def countdown(text: str, t: int):
    while t >= 0:
        print(f"{text} : {t} sec", end="\r")
        sleep(1)
        t -= 1
    print()


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(
        os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


def notary_email():
    clear_display()
    print("\n")
    print_center(
        f"-------------------  Account : {user.email}  -------------------")
    print()
    print_center("Notary Email")
    print("\n")
    while True:
        url = input("Target Google Sheet Link ( 0 : quit ) : ").strip()
        if url == "0":
            return
        try:
            spreadsheet = gc.open_by_url(url)
            break
        except SpreadsheetNotFound:
            print("The specified spreadsheet was not found.")
        except HttpError:
            print("\nNo Internet Connection\n")
        except Exception as e:
            print(f"ERROR")
    clear_display()
    print("\n")
    print_center(
        f"-------------------  Account : {user.email}  -------------------")
    print()
    print_center("Notary Email")
    print()
    print_center(f"Google Sheet : {spreadsheet.title}")
    print("\n")
    print("1. Send Emails")
    print("2. Change Google Sheet")
    print("q. Main menu")
    print("\nEnter your choice (1/2/q): ")
    while True:
        if msvcrt.kbhit():
            choice = msvcrt.getch().decode('utf-8').lower()
            if choice == "1":
                print("\nLoading...")
                send_notary_emails(spreadsheet)
                break
            elif choice == "2":
                notary_email()
                break
            elif choice == "q":
                main()
                break
        sleep(0.1)
    while msvcrt.kbhit():
        msvcrt.getch()
    clear_display()
    print("\n")
    print_center(
        f"-------------------  Account : {user.email}  -------------------")
    print()
    print_center("Notary Email")
    print()
    input("\n\nTask Completed\nPress Enter To Continue : ")
    main()


def client_email():
    clear_display()
    print("\n")
    print_center(f"-------------------  Account : {user.email}  -------------------")
    print()
    print_center("Client Email")
    print()
    user_input = input(f"Enter a list of rows separated by commas ( 0 : quit ) : ")
    input_list = user_input.split(",")
    spreadsheet = gc.open_by_key(INVOICE_SHEET_KEY)
    worksheet = spreadsheet.get_worksheet(0)
    row_value = worksheet.row_values(row)
    primary_header_row = get_filled_row_values(worksheet,4)
    secondary_header_row = get_filled_row_values(worksheet,5)
    for item in input_list:
        item = item.strip()
        try:
            row = int(item)
        except:
            print("Invalid input. Please enter only integers separated by commas.")
            continue
        if row <= 5:
            return
        try:
            print(f"\n\nCreating Draft for row {row}")
            person_full_name = row_value[find_column_index(primary_header_row, secondary_header_row,"Nom/Prénom")]
            amount_found_by_us = row_value[find_column_index(primary_header_row, secondary_header_row,"Somme retrouvée")]
            amount_with_tex = row_value[find_column_index(primary_header_row, secondary_header_row,"Commission TTC (notaire déj payé)")]
            amount_after_fee = row_value[find_column_index(primary_header_row, secondary_header_row,"Somme à verser (incl cas spécifique EON)")]
            message = create_client_message(user.email, "", person_full_name, amount_found_by_us, amount_with_tex, amount_after_fee)
            status = create_draft(message)
            if status:
                input("\nSuccess    ")
            else:
                input("\nError      ")
            client_email()
        except Exception as e:
            print(f"ERROR : {e}")
    input("\nPress Enter to Continue :")
    client_email()


def facturation():
    clear_display()
    print("\n")
    print_center(
        f"-------------------  Account : {user.email}  -------------------")
    print()
    print_center("Facturation")
    print()
    user_input = input(f"Enter a list of rows separated by commas ( 0 : quit ) : ")
    input_list = user_input.split(",")
    spreadsheet = gc.open_by_key(INVOICE_SHEET_KEY)
    worksheet = spreadsheet.get_worksheet(0)
    primary_header_row = get_filled_row_values(worksheet,4)
    secondary_header_row = get_filled_row_values(worksheet,5)
    for item in input_list:
        item = item.strip()
        try:
            row = int(item)
        except:
            print("Invalid input. Please enter only integers separated by commas.")
            continue
        if row <= 5:
            return
        try:
            print(f"\n\nCreating Draft for row {row}")
            row_value = worksheet.row_values(row)
            person_full_name = row_value[find_column_index(primary_header_row, secondary_header_row,"Nom/Prénom")]
            facture_number = row_value[find_column_index(primary_header_row, secondary_header_row,"# Factures LD","LD")]
            ht = row_value[find_column_index(primary_header_row, secondary_header_row,"Commission HT","LD")]
            tva = row_value[find_column_index(primary_header_row, secondary_header_row,"TVA Commission","LD")]
            tcc = row_value[find_column_index(primary_header_row, secondary_header_row,"Commission TTC","LD")]
            try:
                paid_date = datetime.strptime(row_value[find_column_index(primary_header_row, secondary_header_row,"Date paiement","LD")], '%d/%m/%Y').strftime('%d %B %Y')
            except:
                print("No Paiement Date")
                paid_date = ""
            message = create_facture_message(user.email, "", person_full_name)
            status = create_draft(message)
            if status:
                print(f"Creating Invoice for row {row}")
                create_facture_files(person_full_name, facture_number, ht, tva, tcc, paid_date)
                print(f"{row} Success")
            else:
                print(f"{row} Error")
        except Exception as e:
            print(f"{row} ERROR : {e}")
    input("\n\nPress Enter to Continue :")
    facturation()


def main():
    clear_display()
    print("\n")
    print_center(
        f"-------------------  Account : {user.email}  -------------------")
    print("\n")
    print("1. Notary Email")
    print("2. Client Email")
    print("3. Facturation")
    print("\nEnter your choice (1/2/3): ")
    while True:
        if msvcrt.kbhit():
            choice = msvcrt.getch().decode('utf-8').lower()
            if choice == "1":
                print("\nLoading...")
                notary_email()
                break
            elif choice == "2":
                print("\nLoading...")
                client_email()
                break
            elif choice == "3":
                print("\nLoading...")
                facturation()
                break
        sleep(0.1)
    while msvcrt.kbhit():
        msvcrt.getch()
    main()


load_dotenv(dotenv_path=resource_path(".env"))
client_secret = os.environ["CLIENT_SECRET"]
client_secret_info = json.loads(client_secret)

 
NOTARY_SHEET_KEY = "1NBWDbmuXHKr6yWsEvxJhio4uaUPKol6_dJvtgKJCDhc"
INVOICE_SHEET_KEY = "1KlKBSzyFDprXy_L8Gy0UDfRfMdmpl-YZnZErg0yiATg"
locale.setlocale(locale.LC_TIME, 'fr_FR')
if __name__ == "__main__":
    try:
        check_for_updates()
        print("Running the latest version.")
        user = GoogleServices()
        gc = gspread.authorize(user.creds)
        notary_sheet = gc.open_by_key(NOTARY_SHEET_KEY)
        notary_worksheet = notary_sheet.get_worksheet(0)
        scheduling_worksheet = notary_sheet.get_worksheet_by_id(1111177424)   
        main()
    except Exception as e:
        print(e)
        print("\n\n!! Error !!")
    input("Press Enter to EXIT : ")
